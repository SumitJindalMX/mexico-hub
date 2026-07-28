"""Generate Amdocs Mexico Hub end-user Word guide."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "docs" / "Mexico-Hub-User-Guide.docx"

# Mexico-ish accents
GREEN = RGBColor(0x00, 0x6B, 0x34)
RED = RGBColor(0xCE, 0x11, 0x26)
DARK = RGBColor(0x1F, 0x2A, 0x24)
MUTED = RGBColor(0x4A, 0x5A, 0x52)


def set_run_font(run, size=11, bold=False, italic=False, color=DARK, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_heading_styled(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        if level == 1:
            set_run_font(run, size=18, bold=True, color=GREEN, name="Calibri")
        elif level == 2:
            set_run_font(run, size=14, bold=True, color=RED, name="Calibri")
        else:
            set_run_font(run, size=12, bold=True, color=DARK, name="Calibri")
    return p


def add_para(doc, text, *, bold=False, italic=False, size=11, color=DARK, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11, color=DARK)
        p.paragraph_format.space_after = Pt(3)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run, size=11, color=DARK)
        p.paragraph_format.space_after = Pt(3)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "006B34")
        shading.set(qn("w:val"), "clear")
        hdr[i]._te = hdr[i]._tc.get_or_add_tcPr()
        hdr[i]._tc.get_or_add_tcPr().append(shading)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(val)
            set_run_font(run, size=10, color=DARK)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Amdocs Mexico Hub")
    set_run_font(r, size=28, bold=True, color=GREEN, name="Calibri")

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run("make it amazing")
    set_run_font(r, size=16, italic=True, color=RED, name="Calibri")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("User Guide — Creating Activities & Participant Registration")
    set_run_font(r, size=12, color=MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "Live site: https://sumitjindalmx.github.io/mexico-hub/\n"
        "Document version: 1.0 · July 2026"
    )
    set_run_font(r, size=10, color=MUTED)

    doc.add_paragraph()

    # 1 Overview
    add_heading_styled(doc, "1. Overview", 1)
    add_para(
        doc,
        "Amdocs Mexico Hub is the public platform for Mexico activities — hackathons, "
        "culture, talent, sports, leadership visits, and site programs. Editors publish "
        "activities; participants register teams and optionally attach PPT/video materials.",
    )
    add_bullets(
        doc,
        [
            "Browse and filter activities by category, status, visibility, and confidence",
            "Editors create and publish new activities (GitHub allowlist)",
            "Participants register teams (Google, Gmail, GitHub Issue, or Microsoft when available)",
            "Optional PPT/video via Google Drive upload or pasted links",
            "Organizers can generate invite codes and email them via Gmail",
        ],
    )

    # 2 Roles
    add_heading_styled(doc, "2. Who does what", 1)
    add_table(
        doc,
        ["Role", "What they do", "Sign-in needed"],
        [
            [
                "Visitor",
                "Browse activities, open details, filter the catalog",
                "None",
            ],
            [
                "Participant",
                "Register a team, optional invite code, optional PPT/video",
                "Google recommended (for Drive upload); or submit via Gmail / GitHub Issue",
            ],
            [
                "Editor",
                "Create / publish activities to the live catalog",
                "Editor (GitHub PAT) — allowlisted GitHub username",
            ],
            [
                "Organizer",
                "Generate invite codes, email invites, review registrations",
                "Editor and/or Google; Microsoft when SharePoint consent is granted",
            ],
        ],
    )

    # 3 Browse
    add_heading_styled(doc, "3. Browse activities", 1)
    add_numbered(
        doc,
        [
            "Open https://sumitjindalmx.github.io/mexico-hub/ (use the trailing slash).",
            "Hard-refresh (Ctrl+Shift+R) after updates if the page looks stale.",
            "Scroll to Activities. Use filters: Category, Status, Visibility, Confidence, and search.",
            "Click an activity in the list to see details on the right.",
            "If registration is open, use Register team & upload PPT/video.",
            "Use Manage invites to create/share invite codes (organizers / editors).",
        ],
    )
    add_para(
        doc,
        "Confidence badges: Verified (ops-confirmed), Editor (published by allowlisted editor), "
        "Seed (illustrative public-source seed — not an official calendar).",
        italic=True,
        color=MUTED,
        size=10,
    )

    # 4 Create event
    add_heading_styled(doc, "4. Create an activity (Editors)", 1)
    add_heading_styled(doc, "4.1 Editor sign-in", 2)
    add_numbered(
        doc,
        [
            "Ask a site admin to add your GitHub username to js/auth-config.js (authorizedUsers) and grant write access to the mexico-hub repo.",
            "Create a GitHub Personal Access Token (fine-grained) with Contents: Read and write on SumitJindalMX/mexico-hub.",
            "On the site, click Editor (top bar).",
            "Paste the PAT and sign in. The token stays in this browser tab only (sessionStorage).",
            "Create activity appears when you are signed in as an editor.",
        ],
    )
    add_para(
        doc,
        "Never commit PATs, client secrets, or tokens to the repository.",
        bold=True,
        color=RED,
    )

    add_heading_styled(doc, "4.2 Publish an activity", 2)
    add_numbered(
        doc,
        [
            "Click Create activity.",
            "Fill required fields (marked with a red *): Name, Category, Status, Visibility, Confidence, When, Audience, Highlight.",
            "Optional: Sort key (date), PPT/deck URL, Video URL.",
            "Optional: check Open participant registration so teams can register.",
            "Click Publish event. GitHub Pages usually updates within about one minute.",
        ],
    )
    add_para(
        doc,
        "City labels such as GDL mean Guadalajara when a location is needed. Brand the activity as Mexico-first.",
        italic=True,
        color=MUTED,
        size=10,
    )

    # 5 Register
    add_heading_styled(doc, "5. Participant registration", 1)
    add_heading_styled(doc, "5.1 Required vs optional", 2)
    add_table(
        doc,
        ["Field", "Required?"],
        [
            ["Team name", "Yes *"],
            ["Lead name", "Yes *"],
            ["Lead email", "Yes *"],
            ["Team members (at least one name)", "Yes *"],
            ["Invite code", "No (optional)"],
            ["PPT file / PPT URL", "No (optional)"],
            ["Video file / Video URL", "No (optional)"],
        ],
    )

    add_heading_styled(doc, "5.2 How to register a team", 2)
    add_numbered(
        doc,
        [
            "Open an activity that has registration open.",
            "Click Register team & upload PPT/video.",
            "(Recommended) Click Google in the top bar and sign in with your Google account.",
            "Enter team name, lead name, lead email, and at least one member name.",
            "Optionally enter an invite code from the organizer.",
            "Optionally attach PPT/video files (Google Drive) or paste links.",
            "Submit using one of the options below.",
        ],
    )

    add_heading_styled(doc, "5.3 Submit options", 2)
    add_table(
        doc,
        ["Button / path", "When to use", "What happens"],
        [
            [
                "Submit registration (Google signed in)",
                "Default for participants",
                "If files were chosen, uploads to your Google Drive, then opens Gmail compose to the organizer with registration details + downloads JSON",
            ],
            [
                "Submit registration (Editor signed in)",
                "Editors publishing a registration into the repo",
                "Saves into data/registrations.json on GitHub",
            ],
            [
                "Submit via Gmail",
                "No editor session; works without Microsoft",
                "Opens Gmail compose to the registration inbox; downloads a JSON copy",
            ],
            [
                "Submit via GitHub Issue",
                "Track registrations as GitHub issues",
                "Opens a prefilled issue on mexico-hub; downloads JSON",
            ],
            [
                "Microsoft / SharePoint",
                "Only when Entra admin consent is granted",
                "Writes to SharePoint lists and uploads library",
            ],
        ],
    )

    add_heading_styled(doc, "5.4 PPT and video materials (optional)", 2)
    add_bullets(
        doc,
        [
            "You can register without PPT or video.",
            "With Google sign-in: choose files → they upload to your Google Drive and become shareable links on the registration.",
            "Without Google: paste PPT/video URLs (Drive, SharePoint, Stream, etc.).",
            "Size guidance: ~50 MB PPT, ~250 MB video.",
            "First-time Google: enable Google Drive API in Google Cloud and approve Drive access when prompted.",
        ],
    )

    # 6 Invites
    add_heading_styled(doc, "6. Invite codes (Organizers)", 1)
    add_numbered(
        doc,
        [
            "Open an activity → Manage invites.",
            "Set Max uses and optional Expiry.",
            "Optional: enter emails to receive the invite.",
            "Generate & email via Gmail — creates a code and opens Gmail.",
            "Generate & publish to GitHub — saves codes to data/invites.json (requires Editor sign-in).",
            "Share the code with participants; they enter it on Register team (optional field).",
        ],
    )

    # 7 Google setup
    add_heading_styled(doc, "7. Google sign-in notes", 1)
    add_para(
        doc,
        "If Google shows Error 403 access_denied / “app is being tested”:",
    )
    add_numbered(
        doc,
        [
            "Open https://console.cloud.google.com/auth/audience (new Google Auth Platform → Audience).",
            "Select the project that owns the Mexico Hub OAuth client.",
            "Under Test users, add the exact Gmail addresses that should sign in — or Publish app for everyone.",
            "Hard-refresh the site and try Google again.",
        ],
    )
    add_para(
        doc,
        "Full Google setup (OAuth client, redirect URI, Drive API): see google/setup.md in the repository.",
        italic=True,
        color=MUTED,
        size=10,
    )

    # 8 Microsoft
    add_heading_styled(doc, "8. Microsoft / SharePoint (optional)", 1)
    add_para(
        doc,
        "Microsoft sign-in and SharePoint list/file storage need Amdocs Entra admin consent. "
        "Until that is granted, use Google Drive uploads and Gmail / GitHub Issue submission paths. "
        "Details: sharepoint/lists-setup.md in the repository.",
    )

    # 9 Troubleshooting
    add_heading_styled(doc, "9. Troubleshooting", 1)
    add_table(
        doc,
        ["Problem", "What to try"],
        [
            [
                "Page CSS/JS missing",
                "Open the URL with a trailing slash: …/mexico-hub/",
            ],
            [
                "Google 403 access_denied",
                "Add your email under Audience → Test users, or Publish app",
            ],
            [
                "Cannot upload files",
                "Sign in with Google (Drive), or paste links instead",
            ],
            [
                "Editor not allowed",
                "Ask admin to add your GitHub login to auth-config allowlist",
            ],
            [
                "Registration not visible on site",
                "Gmail/Issue paths are manual; Editor path updates data/registrations.json after Pages rebuild (~1 min)",
            ],
            [
                "Microsoft sign-in pending approval",
                "Expected without Entra admin consent — use Google / Gmail / GitHub",
            ],
        ],
    )

    # 10 Quick checklist
    add_heading_styled(doc, "10. Quick checklists", 1)
    add_heading_styled(doc, "Editor — publish activity", 3)
    add_bullets(
        doc,
        [
            "□ Allowlisted on GitHub + PAT ready",
            "□ Editor sign-in successful",
            "□ Required fields filled (*)",
            "□ Registration open checked if teams should register",
            "□ Publish event → wait for Pages refresh",
        ],
    )
    add_heading_styled(doc, "Participant — register team", 3)
    add_bullets(
        doc,
        [
            "□ Activity has registration open",
            "□ Google sign-in (recommended)",
            "□ Team + lead + member name filled (*)",
            "□ PPT/video optional",
            "□ Submit via registration / Gmail / GitHub Issue",
        ],
    )

    add_para(doc, "")
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = end.add_run("Amdocs · make it amazing · Mexico Hub")
    set_run_font(r, size=11, italic=True, color=GREEN)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
